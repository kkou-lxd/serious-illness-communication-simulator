"""Helpers for extracting text from uploaded case files."""

from __future__ import annotations

import io
import re
import tempfile
from pathlib import Path
from typing import Callable, List, Optional

import pdfplumber
import pytesseract
from docx import Document
from pdf2image import convert_from_bytes
from PIL import Image
from streamlit.runtime.uploaded_file_manager import UploadedFile


def _normalize_text(text: str) -> str:
    return "\n".join(line.strip() for line in text.replace("\r", "").splitlines()).strip()


SECTION_STOP_RE = re.compile(r"^\s*(CLINICIAN|PARENT)\b", re.IGNORECASE)
GENERAL_DESC_RE = re.compile(r"^\s*general\s*description\s*:?\s*(.*)$", re.IGNORECASE)


def _extract_txt(file: UploadedFile) -> str:
    return _normalize_text(file.read().decode("utf-8", errors="ignore"))


def _extract_docx(file: UploadedFile) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file.read())
        tmp_path = Path(tmp.name)
    try:
        doc = Document(tmp_path)
        paragraphs: List[str] = [p.text for p in doc.paragraphs if p.text]
        # Include table text rows
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells if cell.text]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return _normalize_text("\n".join(paragraphs))
    finally:
        tmp_path.unlink(missing_ok=True)


def _extract_general_description_from_text(text: str) -> str:
    lines = [line.strip() for line in (text or "").replace("\r", "").splitlines()]
    start_index: Optional[int] = None
    collected: List[str] = []
    for idx, line in enumerate(lines):
        match = GENERAL_DESC_RE.match(line)
        if match:
            start_index = idx
            inline = (match.group(1) or "").strip()
            if inline:
                collected.append(inline)
            break
    if start_index is None:
        return ""
    for line in lines[start_index + 1 :]:
        if SECTION_STOP_RE.match(line):
            break
        if line:
            collected.append(line)
    return _normalize_text("\n".join(collected))


def _extract_general_description_docx(file: UploadedFile) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file.read())
        tmp_path = Path(tmp.name)
    try:
        doc = Document(tmp_path)
        # Prefer table-based section extraction first.
        for table in doc.tables:
            rows = table.rows
            for row_idx, row in enumerate(rows):
                cells = [cell.text.strip() for cell in row.cells]
                for col_idx, cell_text in enumerate(cells):
                    match = GENERAL_DESC_RE.match(cell_text)
                    if not match:
                        continue
                    inline = (match.group(1) or "").strip()
                    if inline:
                        return _normalize_text(inline)
                    if col_idx + 1 < len(cells) and cells[col_idx + 1]:
                        return _normalize_text(cells[col_idx + 1])
                    collected: List[str] = []
                    for next_row in rows[row_idx + 1 :]:
                        row_texts = [c.text.strip() for c in next_row.cells if c.text.strip()]
                        if not row_texts:
                            continue
                        if any(SECTION_STOP_RE.match(t) for t in row_texts):
                            break
                        collected.append(" ".join(row_texts))
                    if collected:
                        return _normalize_text("\n".join(collected))
        # Fallback to paragraph/text section extraction.
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        full_text = _normalize_text("\n".join(paragraphs))
        return _extract_general_description_from_text(full_text)
    finally:
        tmp_path.unlink(missing_ok=True)


def _extract_pdf_with_text(file_bytes: bytes) -> str:
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        pages_text = []
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            # Preserve intra-page line breaks
            pages_text.append(page_text.rstrip())
    extracted = "\n\n".join(pages_text)
    return _normalize_text(extracted)


def _extract_pdf_with_ocr(file_bytes: bytes) -> str:
    images = convert_from_bytes(file_bytes)
    text_chunks = []
    for img in images:
        gray = img.convert("L")
        text_chunks.append(pytesseract.image_to_string(gray))
    return _normalize_text("\n".join(text_chunks))


def _extract_pdf(file: UploadedFile) -> str:
    file_bytes = file.read()
    text_direct = _extract_pdf_with_text(file_bytes)
    if text_direct:
        return text_direct
    return _extract_pdf_with_ocr(file_bytes)


def _fallback_pypdf2_extract(file_bytes: bytes) -> str:
    try:
        import PyPDF2  # type: ignore
    except ImportError:
        return ""
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    pages_text = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        pages_text.append(page_text.rstrip())
    return _normalize_text("\n\n".join(pages_text))


_HANDLERS: dict[str, Callable[[UploadedFile], str]] = {
    "txt": _extract_txt,
    "docx": _extract_docx,
    "pdf": _extract_pdf,
}


def extract_text_from_uploaded_file(file: UploadedFile) -> str:
    """
    Dispatch text extraction for supported file types.
    Returns normalized plain text; raises ValueError on unsupported formats.
    """

    suffix = (Path(file.name).suffix or "").lower().lstrip(".")
    handler = _HANDLERS.get(suffix)
    if not handler:
        raise ValueError("Unsupported file type. Please upload .txt, .pdf, or .docx")
    primary = handler(file)
    # Fallback for suspiciously short PDF extraction
    if suffix == "pdf" and len(primary) < 200:
        file.seek(0)
        fallback = _fallback_pypdf2_extract(file.read())
        if len(fallback) > len(primary):
            return fallback
    return primary


def extract_general_description_from_uploaded_file(file: UploadedFile) -> str:
    """Extract only the 'General Description' section from an uploaded file."""
    suffix = (Path(file.name).suffix or "").lower().lstrip(".")
    if suffix == "docx":
        return _extract_general_description_docx(file)
    # For txt/pdf, do best-effort section slicing from full extracted text.
    full_text = extract_text_from_uploaded_file(file)
    return _extract_general_description_from_text(full_text)


def extract_text_from_path(path: Path) -> str:
    """
    Extract text from a path on disk using the same rules as uploaded files.
    """
    suffix = path.suffix.lower().lstrip(".")
    fake_upload = type(
        "FileShim",
        (),
        {
            "name": path.name,
            "read": lambda self: path.read_bytes(),
            "seek": lambda self, x: None,
        },
    )()
    handler = _HANDLERS.get(suffix)
    if not handler:
        if suffix in {"txt"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        raise ValueError(f"Unsupported file type: {suffix}")
    return handler(fake_upload)  # type: ignore


__all__ = ["extract_text_from_uploaded_file", "extract_general_description_from_uploaded_file"]
